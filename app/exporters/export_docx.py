from __future__ import annotations

from pathlib import Path

from app.database.repository import Repository
from app.exporters.validation import chapter_text, validate_export_chapters
from app.exporters.naming import book_export_path, chapter_export_path, chapter_range_export_path


def export_docx(
    repo: Repository,
    work_id: int,
    output_path: Path | None = None,
    *,
    include_draft: bool = False,
) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx，请先运行 pip install -r requirements.txt") from exc

    work = repo.get_work(work_id)
    chapters = repo.chapters_for_export(work_id)
    ready_chapters = validate_export_chapters(chapters, include_draft=include_draft)
    output_path = output_path or book_export_path(work, "docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(work.get("title") or "未命名作品", level=0)
    if work.get("summary"):
        doc.add_heading("简介", level=1)
        doc.add_paragraph(work["summary"])

    for chapter, text in ready_chapters:
        title = chapter.get("title") or f"第{chapter['chapter_number']}章"
        doc.add_heading(title, level=1)
        for paragraph in text.splitlines():
            paragraph = paragraph.strip()
            if paragraph:
                doc.add_paragraph(paragraph)

    doc.save(output_path)
    return output_path


def export_chapter_docx(
    repo: Repository,
    work_id: int,
    chapter_number: int,
    output_path: Path | None = None,
    *,
    include_draft: bool = False,
) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx，请先运行 pip install -r requirements.txt") from exc

    work = repo.get_work(work_id)
    chapter = repo.get_chapter(work_id, chapter_number)
    output_path = output_path or chapter_export_path(work, chapter, "docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(work.get("title") or "未命名作品", level=0)
    doc.add_heading(chapter.get("title") or f"第{chapter_number}章", level=1)
    text = chapter_text(chapter, include_draft=include_draft)
    if not text.strip():
        raise ValueError("当前章节没有最终稿，不能导出正式稿。")
    for paragraph in text.splitlines():
        paragraph = paragraph.strip()
        if paragraph:
            doc.add_paragraph(paragraph)

    doc.save(output_path)
    return output_path


def export_range_docx(
    repo: Repository,
    work_id: int,
    start_chapter: int,
    end_chapter: int,
    output_path: Path | None = None,
    *,
    include_draft: bool = False,
) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx，请先运行 pip install -r requirements.txt") from exc

    work = repo.get_work(work_id)
    chapters = [
        chapter
        for chapter in repo.chapters_for_export(work_id)
        if start_chapter <= int(chapter.get("chapter_number") or 0) <= end_chapter
    ]
    ready_chapters = validate_export_chapters(
        chapters,
        include_draft=include_draft,
        start=start_chapter,
        end=end_chapter,
    )
    output_path = output_path or chapter_range_export_path(work, start_chapter, end_chapter, "docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(work.get("title") or "未命名作品", level=0)
    doc.add_paragraph(f"第{start_chapter:03d}-{end_chapter:03d}章")
    for chapter, text in ready_chapters:
        title = chapter.get("title") or f"第{chapter['chapter_number']}章"
        doc.add_heading(title, level=1)
        for paragraph in text.splitlines():
            paragraph = paragraph.strip()
            if paragraph:
                doc.add_paragraph(paragraph)

    doc.save(output_path)
    return output_path
